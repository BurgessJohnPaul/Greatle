import json
import csv
from docx import Document

MAX_CHARACTERS = 49000

def isEnglish(s):
    try:
        s.encode(encoding='utf-8').decode('ascii')
    except UnicodeDecodeError:
        return False
    else:
        return True

def delimited_quote(quote, tags=None, author=None):
    if not isEnglish(quote):
        return ''
    d_tags = ''
    if tags is not None:
        tags = [tag.strip() for tag in tags]
        d_tags = '#{};'.format(';'.join(tags))
    d_author = '' if author is None else '@{}'.format(author.split(',')[0].strip())
    return '{}{}^{}^\n'.format(d_tags, d_author, quote.strip())

def parse_quotes_json_object(obj):
    tags = obj["Tags"] if "Tags" in obj else None
    author = obj["Author"] if "Author" in obj else None
    return delimited_quote(obj["Quote"], tags, author)

def save_quotes_to_word(quotes, name):
    numCharacters = 0
    docNum = 0
    document = Document()
    for quote in quotes:
        document.add_paragraph(quote)
        numCharacters += len(quote)
        if numCharacters >= MAX_CHARACTERS:
            document.save(name + str(docNum) + '.docx')
            document = Document()
            numCharacters = 0
            docNum += 1
    document.save(name + str(docNum) + '.docx')

quotes = set()
with open('quotes.json', encoding='windows-1256') as f:
    data = json.load(f)
for obj in data:
    quotes.add(parse_quotes_json_object(obj))
quotes.remove('')
save_quotes_to_word(quotes, 'quotes')

quotes.clear()
with open('quotes_all.csv') as csv_file:
    csv_reader = csv.reader(csv_file, delimiter=';')
    for row in csv_reader:
        quotes.add(delimited_quote(row[0], author=row[1]))
save_quotes_to_word(quotes, 'quotes_all')
